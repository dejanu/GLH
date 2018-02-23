#https://automatetheboringstuff.com/chapter13/
import docx
d=docx.Document('django.docx')
d.paragraphs[0].text
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
